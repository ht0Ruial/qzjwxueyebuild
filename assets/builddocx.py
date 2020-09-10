from docx import Document
from docxtpl import DocxTemplate
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from os import remove
from os.path import dirname
from os.path import abspath
import time
from getinfo import list_info, list_mo_info, last_course, now_course



# 宋体五号

def fontsize(run):
    run.font.size = Pt(10.5)


# 字体加粗

def adbold(run):
    run.font.bold = True


# 处理拼接问题

def do_pj(i, j, list,color=False):
    cell = table.cell(i, j)
    run = cell.paragraphs[0].add_run(list)
    run.font.size = Pt(12)  # 宋体小四
    if i in (1, 11, 12):
        fontsize(run)  # 宋体五号
    if i in (1, 11):
        adbold(run)  # 加粗
    if color == True:
        run.font.color.rgb = RGBColor(255, 0, 0)


# 新增单元格合并

def addrow_merge(table, i, list_te):
    table.add_row().cells
    for j in range(len(list_te)):
        cell_1 = table.cell(i, list_te[j][0])
        cell_2 = table.cell(i, list_te[j][1])
        cell_1.merge(cell_2)


# 课程写入表中

def do_cell_center(table, i, j, list, k):
    cell = table.cell(i, j)
    run = cell.paragraphs[0].add_run(list[k])
    run.font.size = Pt(12)  # 宋体小四
    cell.paragraphs[0].paragraph_format.line_spacing = Pt(15)
    cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


data = {
    'grade': list_info[0],
    'studentid': list_info[1],
    'name': list_info[2],
    'sex': list_info[3],
    'cc': list_info[4],
    'xy': list_info[5],
    'major': list_info[6],
}

document = DocxTemplate('.\\assets\\demo.docx')
document.render(data)
document.save('.\\assets\\test.docx')

document = Document('.\\assets\\test.docx')

p = document.add_paragraph()

table = document.tables[0]

list_b = [
    14,  # 0
    14,  # 1
    0, 2, 4, 6, 8,  # GPA
    0, 8, 14,
    0, 8, 14,  # 4体测
    0, 8,
    0, 8, 14,  # 6 四级
    0, 8,
    0, 8,  # 8专业其他
    0,  # 9科技活动
    0, 14,  # 10学籍
    0, 3, 5, 7, 8, 10, 11, 12, 13, 14,  # 13课程
    0  # 防溢出
]
list_c = ['学生毕业要求及本人情况对照', '其他毕业要求及学籍异动', '课程类别', '学分要求', '取得学分',
          '欠修学分', 'GPA要求', '公共必修课程', '综合素质测评', '要求：各学期综合素质测评成绩合格。目前是否都合格（ 合格 ）',
          '公共选修课程', '体质健康', '要求：体质测试成绩合格。目前情况：  合格            ',
          '专业必修课程', '双语修课', '专业选修课程', '英语水平', '达到英语水平标准： 大学英语四级   目前是否达到标准（  是 ）',
          '综合实践课程', '专业必修课GPA要求', '创业就业导向课程', '专业其它要求', '课外科技活动', '',
          '本人学籍异动情况：1.（√）无；2.本人于20X 年X月曾休学X年，现复学到20X 级XX专业；3.或本人于20X年X月从XX专业转至XX专业。（直接在对应数字上对“√”并填写相关内容）',
          '课程类型', '课程代码', '课程名称', '成绩等级', '取得学分', '课程类型', '课程代码', '课程名称', '学分', '备注（初修/重修/辅修其它专业课程',
          ]
k = -1
lx_num = 8 
for i in (0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 13):  # 从0到学籍10那一行,再加个13

    if i > 2 and i < 10:
        for l in (2, 4, 6):
            cell = table.cell(i, l)
            run = cell.paragraphs[0].add_run(list_info[lx_num])
            lx_num +=1

    for j in list_b[k+1:]:
        k += 1
        cell = table.cell(i, j)
        run = cell.paragraphs[0].add_run(list_c[k])
        run.font.size = Pt(12)  # 宋体小四

        if i in (0, 1, 10):
            fontsize(run)  # 宋体五号

        if i in range(2):
            adbold(run)  # 加粗

        if i > 2 and j == 14:
            # 左对齐
            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        else:
            # 居中
            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        if list_b[k] >= list_b[k+1]:
            break

# 拼接部分
list_b.clear()
list_c.clear()
list_b = [
    5, 6, 13,
    13,
    11, 13, 14,  # 双语
    13,  # 专必GPA
    6, 7, 8, 10, 12, 13,  # 学期
    6, 7, 8, 13, 14,
    0  # 防溢出
]
list_c = [
    '毕业总学分要求≥  ','  ，目前已取得总学分：', '',
    '要求：GPA ≥ 2.00     目前GPA：',
    '要求修读门数：','   已修读门数：', '     欠修门数：',
    '学位要求：GPA ≥ 2.20        目前GPA：',
    '一、上学期即', '学年第', '学期学习执行情况', '二、本学期即', '学年第', '学期修读的课程',
    '取得总学分：', '  ；截止到目前GPA（取得学分）：', '  ；截止到目前GPA（所有课程）：',
    '预期本学期取得学分：', '    预期截止到本学期末GPA（取得学分）：',
    '',
]


k = -1
for i in (1, 2, 5, 7, 11, 12):

    for j in list_b[k+1:]:
        k += 1
        do_pj(i, j, list_c[k])
        if j in (11,14): #上色表格的纵坐标11，14
            color = True
        else:
            color = False
        do_pj(i, j, list_mo_info[k],color)

        if i == 11:
            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        if list_b[k] >= list_b[k+1]:
            break

# 课程部分

if len(last_course)/5 > len(now_course)/4:
    max_len = len(last_course)/5
else:
    max_len = len(now_course)/4

table_len = 13 + int(max_len)  # max_lan是课程的最大行数

a_num, b_num = 0, 0
for i in range(14, table_len+1):
    list_te = [[0, 1], [2, 3], [4, 5], [6, 7], [9, 10]]
    addrow_merge(table, i, list_te)
    
    try:
        for j in (1, 3, 5, 7, 8):
            do_cell_center(table, i, j, last_course, a_num)
            a_num += 1

    except IndexError:
        pass

    try:
        for k in range(10, 14):
            do_cell_center(table, i, k, now_course, b_num)
            b_num += 1
        cell = table.cell(i, 14)
        run = cell.paragraphs[0].add_run('初修') #默认初修
        run.font.size = Pt(12)  # 宋体小四
        cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    except IndexError:
        pass

# 最后两行
list_b.clear()
list_c.clear()
list_b = [
    7, 8, 14,
    3, 4, 13, 14,
    0  # 防溢出
]

now_time = time.strftime('%Y{y}%m{m}%d{d}').format(y='年', m='月', d='日')

ab_info = '本人签名：'+list_info[2]+'      '+now_time

list_c = [
    '三、上学期学分课程之外取得的学业成果', '（包括考研准备、学科竞赛获奖、获资格证书、科研或大创项目结项等）',
    '四、上学期所受处分及累计不合格课程情况（如果没有写“无”）',
    '本人总结（重点写明出现问题、原因分析及后期解决的具体方案）',
    ab_info,
    '学业导师意见（针对学生学习计划的制定及累计不合格课程情况，给出具体的学习指导方案）', '学业导师签名：            年  月  日',
]
k = -1
for i in range(table_len+1, table_len+3):
    list_te = [[0, 8], [9, 14]]
    addrow_merge(table, i, list_te)
    for j in list_b[k+1:]:
        k += 1
        cell = table.cell(i, j)
        if j == 4:
            cell.add_paragraph('').paragraph_format.line_spacing = Pt(15)
            cell.add_paragraph('').paragraph_format.line_spacing = Pt(15)
            p = cell.add_paragraph('')
            run = p.add_run(list_c[k])
            p.paragraph_format.first_line_indent = Pt(
                145)  # 缩进5.11厘米 ,达到缩进12字符的效果
        elif i == table_len+2 and j == 14:
            cell.add_paragraph('').paragraph_format.line_spacing = Pt(15)
            p = cell.add_paragraph('')
            run = p.add_run(list_c[k])
            p.paragraph_format.first_line_indent = Pt(
                145)  # 缩进5.11厘米 ,达到缩进12字符的效果

        else:
            run = cell.paragraphs[0].add_run(list_c[k])

        run.font.size = Pt(12)  # 宋体小四
        p.paragraph_format.line_spacing = Pt(15)  # 行距15磅
        cell.paragraphs[0].paragraph_format.line_spacing = Pt(15)  # 行距15磅

        if j != 8:
            adbold(run)  # 加粗

        if list_b[k] >= list_b[k+1]:
            break

try :
    file_name = list_info[1] + list_info[2] + '学业完成情况表'+'.docx'
    document.save(file_name)
    remove('.\\assets\\test.docx')
    file_pa = abspath(dirname(__file__))
    print('学生学业完成情况表已生成！')
    print('文件保存路径：'+ file_pa[:-7] +'\\'+file_name)
except Exception as e:
    print("[-] %s" % e)

